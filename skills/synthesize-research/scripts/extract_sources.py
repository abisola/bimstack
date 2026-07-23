#!/usr/bin/env python3
"""Extract every research source in a folder into plain-text working files.

Usage:  python3 extract_sources.py <raw-data-folder> <output-folder>

Handles: .docx (paragraphs + tables, headings marked with ##), .xlsx (all
sheets, one pipe-delimited row per line), .csv/.tsv (passed through with a
header note), .sav (SPSS via pyreadstat, value labels applied when present).
Recordings/media and unknown types are listed in the inventory but not parsed.

Writes one .txt per source into <output-folder> plus inventory.md summarising
what was found, sized, and skipped — read inventory.md first.
"""
import csv
import os
import sys


def out_name(path, ext=".txt"):
    base = os.path.splitext(os.path.basename(path))[0]
    return "".join(c if c.isalnum() or c in " -_" else "_" for c in base) + ext


def extract_docx(path):
    import docx
    d = docx.Document(path)
    lines = []
    for p in d.paragraphs:
        if not p.text.strip():
            continue
        prefix = "## " if p.style.name.startswith("Heading") else ""
        lines.append(prefix + p.text)
    for t in d.tables:
        lines.append("=== TABLE ===")
        for r in t.rows:
            lines.append(" | ".join(c.text.strip().replace("\n", " / ") for c in r.cells))
    return "\n".join(lines)


def extract_xlsx(path):
    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=True)
    lines = []
    for ws in wb.worksheets:
        lines.append(f"===== SHEET: {ws.title} ({ws.max_row}x{ws.max_column}) =====")
        for row in ws.iter_rows(values_only=True):
            if any(v is not None and str(v).strip() for v in row):
                lines.append(" | ".join("" if v is None else str(v).strip() for v in row))
    return "\n".join(lines)


def extract_csv(path):
    lines = []
    with open(path, newline="", encoding="utf-8", errors="replace") as f:
        delim = "\t" if path.lower().endswith(".tsv") else ","
        for row in csv.reader(f, delimiter=delim):
            lines.append(" | ".join(c.strip() for c in row))
    return "\n".join(lines)


def extract_sav(path):
    import pyreadstat
    df, meta = pyreadstat.read_sav(path, apply_value_formats=True)
    lines = [f"===== SPSS: {os.path.basename(path)} ({len(df)} rows x {len(df.columns)} cols) ====="]
    lines.append("--- COLUMN LABELS ---")
    for col, label in zip(meta.column_names, meta.column_labels):
        lines.append(f"{col}: {label}")
    lines.append("--- DATA ---")
    lines.append(" | ".join(df.columns))
    for _, row in df.iterrows():
        lines.append(" | ".join("" if v is None else str(v) for v in row))
    return "\n".join(lines)


HANDLERS = {
    ".docx": extract_docx, ".xlsx": extract_xlsx, ".xlsm": extract_xlsx,
    ".csv": extract_csv, ".tsv": extract_csv, ".sav": extract_sav,
}
MEDIA = {".mp3", ".mp4", ".m4a", ".wav", ".mov", ".aac", ".ogg"}


def main():
    src, dst = sys.argv[1], sys.argv[2]
    os.makedirs(dst, exist_ok=True)
    inventory = []
    for root, dirs, files in os.walk(src):
        dirs[:] = [d for d in dirs if not d.startswith(".")]
        for fn in sorted(files):
            if fn.startswith(".") or fn.startswith("~$"):
                continue
            path = os.path.join(root, fn)
            rel = os.path.relpath(path, src)
            ext = os.path.splitext(fn)[1].lower()
            size = os.path.getsize(path)
            if ext in HANDLERS:
                try:
                    text = HANDLERS[ext](path)
                    out = os.path.join(dst, out_name(rel.replace(os.sep, "__")))
                    with open(out, "w") as f:
                        f.write(text)
                    inventory.append((rel, ext, size, f"extracted -> {os.path.basename(out)} ({len(text):,} chars)"))
                except Exception as e:
                    inventory.append((rel, ext, size, f"FAILED: {e}"))
            elif ext in MEDIA:
                inventory.append((rel, ext, size, "media — not parsed (recording; transcript may exist elsewhere)"))
            else:
                inventory.append((rel, ext, size, "skipped (unhandled type)"))
    with open(os.path.join(dst, "inventory.md"), "w") as f:
        f.write("# Source inventory\n\n| file | type | size | status |\n|---|---|---|---|\n")
        for rel, ext, size, status in inventory:
            f.write(f"| {rel} | {ext} | {size:,} | {status} |\n")
    for rel, ext, size, status in inventory:
        print(f"{rel}  [{ext}]  {status}")


if __name__ == "__main__":
    main()
